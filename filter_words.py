import docx
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

def open_docx():
    file=docx.Document('C:\\Users\\User\\Desktop\题库\\2013科普知识试题及答案(选择题300题).docx')
    # print(file.paragraphs[1].text)
    for i in range(1,len(file.paragraphs)):
        # print(file.paragraphs[i].text)
        if file.paragraphs[i].text == '' and file.paragraphs[i] != '答案：':
            print(file.paragraphs[i-1].text)
            print(file.paragraphs[i-2].text)
            print(file.paragraphs[i-3].text)
            print(file.paragraphs[i-4].text)
            # print(file.paragraphs[i].text)



def open_csv():
    file=load_workbook('C:\\Users\\User\\Desktop\\题库\\航天科普答题-试题-07.06(总811道)-(4).xlsx')
    # 获取当前活跃的worksheet,默认就是第一个worksheet
    # ws = wb.active

    # 当然也可以使用下面的方法

    # 获取所有表格(worksheet)的名字
    sheets = file.get_sheet_names()
    # print(sheets)
    # 第一个表格的名称
    sheet_first = sheets[0]
    # 获取特定的worksheet
    ws = file.get_sheet_by_name(sheet_first)

    # 获取表格所有行和列，两者都是可迭代的
    rows = ws.rows
    columns = ws.columns

    # 迭代所有的行
    # for row in rows:
    #     for col in row:
    #         if col.value != None:
    #             print(col.value)

        # print(ws.cell(row=2, column=c).value)
    lines=[]
    for row in range(2,813):

        line=[ws.cell(row=row,column=c).value for c in range(1, 8)]

        lines.append(line)
    print(lines)
    lines=delete_d(lines)
    print(lines)
    save_xlsx(lines)
    # 通过坐标读取值
    # print
    # ws.cell('A1').value  # A表示列,1表示行
    # print
    # ws.cell(row=1, column=1).value

def delete_d(lines):
    for line in lines:
        if line[5] == 'D':
            line[5] = 'C'
            line[3] = line[4]
    return lines

def save_xlsx(lines):
    # 在内存中创建一个workbook对象，而且会至少创建一个 worksheet
    wb = Workbook()

    # 获取当前活跃的worksheet,默认就是第一个worksheet
    ws = wb.active

    # 设置单元格的值，A1等于6(测试可知openpyxl的行和列编号从1开始计算)，B1等于7
    # ws.cell(row=1, column=1).value = 6
    # ws.cell("B1").value = 7

    # 从第2行开始，写入9行10列数据，值为对应的列序号A、B、C、D...
    # for row in range(2, 11):
    #     for col in range(1, 11):
    #         ws.cell(row=row, column=col).value = get_column_letter(col)
    ws.cell(row=1, column=1).value = '题干'
    ws.cell(row=1, column=2).value = '选项A'
    ws.cell(row=1, column=3).value = '选项B'
    ws.cell(row=1, column=4).value = '选项C'
    ws.cell(row=1, column=5).value = '选项D'
    ws.cell(row=1, column=6).value = '正确答案'
    ws.cell(row=1, column=7).value = '难易程度'

    # 可以使用append插入一行数据
    for line in lines:
        ws.append(line)

    # 保存
    wb.save(filename="C:\\Users\\User\\Desktop\\题库\\a.xlsx")

if __name__ == '__main__':
    open_docx()
    # open_csv()